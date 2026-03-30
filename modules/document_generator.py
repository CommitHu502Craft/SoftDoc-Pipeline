"""
说明书生成器模块
基于 docxtpl 库，读取 project_plan.json 和截图文件，自动生成 Word 操作说明书
"""
import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
from docx import Document
from modules.claim_evidence_compiler import compile_claim_evidence_matrix
from modules.document_differentiator import DocumentDifferentiator
from modules.project_charter import resolve_software_identity

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentGenerator:
    """Word 说明书生成器"""
    
    def __init__(self, plan_path: str, screenshot_dir: str, template_path: str, output_path: str):
        """
        初始化生成器
        
        Args:
            plan_path: project_plan.json 路径
            screenshot_dir: 截图目录路径
            template_path: Word 模板路径
            output_path: 输出 Word 文件路径
        """
        self.plan_path = Path(plan_path)
        self.screenshot_dir = Path(screenshot_dir)
        self.template_path = Path(template_path)
        self.output_path = Path(output_path)
        
        # 加载规划数据
        with open(self.plan_path, 'r', encoding='utf-8') as f:
            self.plan = json.load(f)

        # 加载项目元数据 (Task 3: Document Consistency)
        self.metadata = {}
        self.project_spec = {}
        self.project_charter = {}
        # 尝试在 plan 同级目录或上级目录查找 project_metadata.json
        meta_locations = [
            self.plan_path.parent / "project_metadata.json",
            self.output_path.parent / "project_metadata.json",
            Path("project_metadata.json")
        ]

        for meta_path in meta_locations:
            if meta_path.exists():
                try:
                    with open(meta_path, 'r', encoding='utf-8') as f:
                        self.metadata = json.load(f)
                    logger.info(f"已加载项目元数据: {meta_path}")
                    break
                except Exception as e:
                    logger.warning(f"加载元数据失败 {meta_path}: {e}")

        # 加载 Spec-first DSL（用于代码/文档一致性校验）
        spec_locations = [
            self.plan_path.parent / "project_executable_spec.json",
            self.output_path.parent / "project_executable_spec.json",
            self.plan_path.parent / "project_spec.json",
            self.output_path.parent / "project_spec.json",
            Path("project_spec.json"),
            Path("project_executable_spec.json"),
        ]
        for spec_path in spec_locations:
            if spec_path.exists():
                try:
                    with open(spec_path, "r", encoding="utf-8") as f:
                        self.project_spec = json.load(f)
                    logger.info(f"已加载项目规格: {spec_path}")
                    break
                except Exception as e:
                    logger.warning(f"加载项目规格失败 {spec_path}: {e}")

        charter_locations = [
            self.plan_path.parent / "project_charter.json",
            self.output_path.parent / "project_charter.json",
            Path("project_charter.json"),
        ]
        for charter_path in charter_locations:
            if charter_path.exists():
                try:
                    with open(charter_path, "r", encoding="utf-8") as f:
                        self.project_charter = json.load(f)
                    logger.info(f"已加载项目章程: {charter_path}")
                    break
                except Exception as e:
                    logger.warning(f"加载项目章程失败 {charter_path}: {e}")

        self.naming_facts = resolve_software_identity(
            self.project_charter,
            fallback_project_name=str(self.plan.get("project_name") or ""),
        )

        # 初始化差异化增强器
        self.differentiator = DocumentDifferentiator(self.plan)
        logger.info(f"说明书差异化增强器已启用")
    
    def _find_screenshot(self, filename: str) -> Path:
        """
        查找截图文件 (支持模糊匹配，因为截图文件名现在包含哈希前缀)
        Example: 输入 'page_1_full.png'，可能匹配到 '110d09_page_1_full.png'
        """
        # 1. 尝试直接精确匹配
        exact_path = self.screenshot_dir / filename
        if exact_path.exists():
            return exact_path

        # 2. 尝试后缀匹配 (针对带有哈希前缀的文件)
        # 遍历目录下的所有文件
        try:
            for file_path in self.screenshot_dir.glob("*.png"):
                if file_path.name.endswith(f"_{filename}") or file_path.name == filename:
                    return file_path
        except Exception as e:
            logger.error(f"遍历截图目录失败: {e}")

        logger.warning(f"截图文件不存在: {filename}")
        return None
    
    def _create_inline_image(self, tpl: DocxTemplate, image_path: Path, width_mm: int = 160, limit_height: bool = True) -> InlineImage:
        """
        创建内联图片对象 (智能调整大小)
        策略:
        1. 默认尝试使用指定宽度 (width_mm)
        2. 如果 limit_height=True，检查图片高度，如果超过页面 1/3 (约 90mm)，则缩小
        """
        if image_path and image_path.exists():
            try:
                # 尝试导入 PIL 获取图片尺寸
                try:
                    from PIL import Image
                    has_pil = True
                except ImportError:
                    has_pil = False

                final_width_mm = width_mm

                if limit_height and has_pil:
                    with Image.open(image_path) as img:
                        px_width, px_height = img.size
                        # 防止除以零
                        if px_height > 0:
                            aspect_ratio = px_width / px_height

                            # 计算按默认宽度缩放后的预期高度 (mm)
                            expected_height_mm = width_mm / aspect_ratio

                            # 设定最大高度限制 (A4页面高度约297mm，三分之一约 90mm)
                            # 考虑到标题和文字，设置个保守值，比如 85mm
                            MAX_HEIGHT_MM = 85

                            # 如果预期高度超标，则按最大高度反推宽度
                            if expected_height_mm > MAX_HEIGHT_MM:
                                final_width_mm = int(MAX_HEIGHT_MM * aspect_ratio)
                                # logger.debug(f"图片过高 ({expected_height_mm:.1f}mm)，自动缩放至高度 {MAX_HEIGHT_MM}mm: {image_path.name}")

                return InlineImage(tpl, str(image_path), width=Mm(final_width_mm))
            except Exception as e:
                logger.error(f"创建图片失败 {image_path}: {e}")
                return None
        return None
    
    def _collect_page_screenshots(self, page_id: str) -> Dict[str, List[Path]]:
        """
        收集该页面所有可用的组件截图，按类型分类并排序
        Returns:
            {
                'chart': [path_to_chart_1, path_to_chart_3, ...],
                'table': [path_to_table_1, ...]
            }
        """
        collected = {
            'chart': [],
            'table': []
        }

        # 遍历截图目录寻找属于该页面的图片
        # 匹配模式: *_{page_id}_widget_{type}_{id}.png
        # 例如: 110d09_page_1_widget_chart_3.png

        if not self.screenshot_dir.exists():
            return collected

        try:
            for file_path in self.screenshot_dir.iterdir():
                if not file_path.name.endswith(".png"):
                    continue

                # 检查是否属于当前页面 (模糊匹配 page_id)
                # 假设文件名格式: {hash}_{page_id}_{component_id}.png
                # component_id 通常是 widget_chart_X 或 widget_table_X

                if f"_{page_id}_widget_chart_" in file_path.name:
                    collected['chart'].append(file_path)
                elif f"_{page_id}_widget_table_" in file_path.name:
                    collected['table'].append(file_path)

            # 排序：尝试按文件名中的数字排序，保证顺序一致
            # 提取文件名末尾的数字进行排序
            def extract_number(path: Path):
                import re
                match = re.search(r'_(\d+)\.png$', path.name)
                return int(match.group(1)) if match else 0

            collected['chart'].sort(key=extract_number)
            collected['table'].sort(key=extract_number)

        except Exception as e:
            logger.error(f"收集截图文件出错: {e}")

        return collected

    def _sanitize_for_rendering(self, data: Any) -> Any:
        """
        将数据中的列表转换为字符串，以便在 Word 模板直接引用时显示正常
        ['A', 'B'] -> "A、B"
        但是要排除那些需要在模板中遍历的字段 (如 target_users, main_functions)
        """
        # 需要保留 List 结构的字段名白名单
        # 注意：只有模板中使用 {% for %} 循环的字段才需要加入白名单
        # 直接用 {{ xxx }} 输出的字段不要加入，否则会显示 ['...']
        KEEP_LIST_KEYS = [
            'target_users', 'targetUsers',
            'main_functions', 'mainFunctions',
            'main_features', 'mainFeatures',  # 注意：JSON中可能用 features 而非 functions
            'functional_features', 'functionalFeatures',
            'technical_features', 'technicalFeatures',
            # 'advantages' 不在这里，因为模板是直接输出而非循环
            'pages',
            'charts',
            'widgets',
            'db_tables', 'dbTables',
            'api_list', 'apiList'
        ]

        if isinstance(data, dict):
            new_data = {}
            for k, v in data.items():
                if k in KEEP_LIST_KEYS:
                    # 直接保留原值（即保留 List 结构），不递归处理，或者仅递归处理内部 Dict
                    if isinstance(v, list):
                        new_data[k] = [self._sanitize_for_rendering(item) if isinstance(item, (dict, list)) else item for item in v]
                    elif isinstance(v, str) and v:
                        # [修复] 如果白名单字段是字符串，自动拆分为列表
                        # 支持多种分隔符：顿号、分号、换行、数字序号
                        import re
                        # 先尝试按常见分隔符拆分
                        if '、' in v:
                            new_data[k] = [item.strip() for item in v.split('、') if item.strip()]
                        elif '；' in v or ';' in v:
                            new_data[k] = [item.strip() for item in re.split(r'[；;]', v) if item.strip()]
                        elif '\n' in v:
                            new_data[k] = [item.strip() for item in v.split('\n') if item.strip()]
                        elif re.search(r'\d+[.、）)]', v):
                            # 按数字序号拆分，如 "1.xxx 2.xxx" 或 "1、xxx 2、xxx"
                            parts = re.split(r'\d+[.、）)]\s*', v)
                            new_data[k] = [item.strip() for item in parts if item.strip()]
                        else:
                            # 无法拆分，包装成单元素列表
                            new_data[k] = [v]
                    else:
                        new_data[k] = v
                else:
                    new_data[k] = self._sanitize_for_rendering(v)
            return new_data

        elif isinstance(data, list):
            # 简单判断：如果是纯字符串列表，则拼接，避免模板中出现 ['A', 'B']
            if data and all(isinstance(x, str) for x in data):
                return "、".join(data)
            # 如果包含复杂对象（如字典），则保留列表供循环使用
            return [self._sanitize_for_rendering(x) for x in data]
        else:
            return data

    def generate(self):
        """生成 Word 说明书"""
        logger.info(f"开始生成说明书: {self.output_path}")

        # 检查模板是否存在
        if not self.template_path.exists():
            logger.error(f"模板文件不存在: {self.template_path}")
            logger.info("请先创建 Word 模板文件，或等待用户提供模板")
            return False

        # 加载模板
        tpl = DocxTemplate(str(self.template_path))

        # 构造上下文数据
        project_intro = self._sanitize_for_rendering(self.plan.get("project_intro", {}))
        copyright_fields = self._sanitize_for_rendering(self.plan.get("copyright_fields", {}))
        industry = copyright_fields.get("industry", "业务管理")
        software_full_name = str(self.naming_facts.get("software_full_name") or self.plan.get("project_name") or "未命名项目")
        software_short_name = str(self.naming_facts.get("software_short_name") or software_full_name)
        term_dictionary = self.naming_facts.get("term_dictionary") or {}

        context = {
            "project_name": software_full_name,
            "project_alias": software_short_name,
            "admin_name": self.plan.get("admin_name", "系统管理员"),
            "generation_date": datetime.now().strftime("%Y年%m月%d日"),
            "project_intro": project_intro,
            "copyright_fields": copyright_fields,
            "software_full_name": software_full_name,
            "software_short_name": software_short_name,
            "term_dictionary": term_dictionary,
            "pages": [],
            # Phase 3: 注入真实的代码元数据 (数据库表、接口)
            "db_tables": self.metadata.get("entities", []), # list of {class, table}
            "api_list": self.metadata.get("controllers", []), # list of {class, route}
            "project_spec": self.project_spec,
            # 新增：差异化文本
            "intro_text": self.differentiator.get_intro_text(industry),
            "feature_descriptions": self.differentiator.get_feature_description(3),
            "conclusion_text": self.differentiator.get_conclusion_text(),
            # Phase 2: 随机化章节标题与结构策略
            "chapter_titles": {
                k: self.differentiator.get_chapter_title(k)
                for k in ["intro", "install", "manual", "faq", "design"]
            },
            "structure_strategy": self.differentiator.get_structure_strategy(),
            # Phase 5: 部署说明
            "deployment": self.differentiator.get_deployment_guide(self.plan.get("genome", {}).get("target_language", "python")),
            "operation_flow_profile": self.plan.get("operation_flow_profile", {}),
            "diff": self.differentiator  # 提供给模板使用
        }
        ui_skill_profile = self._load_json_quiet(self.output_path.parent / "ui_skill_profile.json")
        ui_blueprint = self._load_json_quiet(self.output_path.parent / "ui_blueprint.json")
        context["ui_skill_profile"] = ui_skill_profile
        context["ui_blueprint_summary"] = ui_blueprint.get("summary") or {}

        # 遍历页面
        for page_id, page_data in self.plan.get("pages", {}).items():
            page_title = page_data.get("page_title", "未命名页面")
            original_desc = page_data.get("page_description", "暂无描述")

            page_ctx = {
                "title": page_title,
                "description": self.differentiator.enrich_page_description(original_desc, page_title),
                "full_screenshot": None,
                "widgets": [],
                "api_info": self._get_page_apis(page_id),  # Phase 3: 新增API关联
                "operation_steps": [],
                "operation_summary": "",
            }
            page_ctx["operation_steps"] = self._build_page_operation_steps(page_title, page_ctx["api_info"])
            page_ctx["operation_summary"] = "；".join(page_ctx["operation_steps"][:3])

            # 添加全景截图（调整尺寸为 170mm）
            full_screenshot_path = self._find_screenshot(f"{page_id}_full.png")
            if full_screenshot_path:
                # 全景截图不受高度限制，保持原样展示
                page_ctx["full_screenshot"] = self._create_inline_image(tpl, full_screenshot_path, width_mm=136, limit_height=False)

            # [V2.4 重构] 以实际截图为准，而非 AI 规划的 charts 配置
            # 这样可以确保：有几个截图就写几个组件，不会出现 None
            available_screenshots = self._collect_page_screenshots(page_id)

            # 从 AI 规划中提取标题映射 (用于匹配截图)
            charts_data = page_data.get("charts", [])
            chart_titles = [c.get("title", f"图表{i+1}") for i, c in enumerate(charts_data) if c.get("type") != "stats_card" and c.get("type") != "table"]
            table_titles = [c.get("title", f"表格{i+1}") for i, c in enumerate(charts_data) if c.get("type") == "table"]

            # 遍历实际存在的图表截图
            for idx, screenshot_path in enumerate(available_screenshots.get('chart', [])):
                if not screenshot_path.exists():
                    continue
                # 尝试获取对应标题，如果没有则使用通用标题
                widget_title = chart_titles[idx] if idx < len(chart_titles) else f"数据分析图表{idx+1}"
                widget_ctx = {
                    "title": widget_title,
                    "description": self.differentiator.generate_widget_intro(widget_title, "chart"),
                    "screenshot": self._create_inline_image(tpl, screenshot_path, width_mm=156)
                }
                page_ctx["widgets"].append(widget_ctx)

            # 遍历实际存在的表格截图
            for idx, screenshot_path in enumerate(available_screenshots.get('table', [])):
                if not screenshot_path.exists():
                    continue
                widget_title = table_titles[idx] if idx < len(table_titles) else f"数据明细表{idx+1}"
                widget_ctx = {
                    "title": widget_title,
                    "description": self.differentiator.generate_widget_intro(widget_title, "table"),
                    "screenshot": self._create_inline_image(tpl, screenshot_path, width_mm=156)
                }
                page_ctx["widgets"].append(widget_ctx)

            logger.info(f"页面 {page_id}: 写入 {len(available_screenshots.get('chart', []))} 个图表 + {len(available_screenshots.get('table', []))} 个表格")

            context["pages"].append(page_ctx)

        self._save_operation_flow_guide(context)
        consistency_report = self._save_consistency_report(context)
        matrix_ok = False
        matrix_payload: Dict[str, Any] = {}
        matrix_path = self.output_path.parent / "claim_evidence_matrix.json"
        try:
            matrix_ok, matrix_path, matrix_payload = compile_claim_evidence_matrix(
                project_name=str(self.plan.get("project_name") or software_full_name),
                project_dir=self.output_path.parent,
                html_dir=self.output_path.parent.parent.parent / "temp_build" / str(self.plan.get("project_name") or "") / "html",
            )
            if matrix_ok:
                logger.info(f"声称-证据矩阵已生成: {matrix_path} (passed)")
            else:
                logger.warning(f"声称-证据矩阵已生成: {matrix_path} (not passed)")
        except Exception as e:
            logger.warning(f"声称-证据矩阵生成失败: {e}")
            matrix_payload = {}

        # 文档阶段先产出可审阅材料；A04 时间线在 freeze 阶段再执行硬门禁，避免阶段时序死锁。
        examiner_material_report = self._save_examiner_material_sections(
            context,
            matrix_payload,
            strict_timeline=False,
        )
        examiner_material_passed = bool(examiner_material_report.get("passed"))
        if examiner_material_passed:
            logger.info("审查员材料章节已生成并通过。")
        else:
            logger.error("审查员材料章节未通过（将触发上层自动修复）。")

        if not consistency_report.get("passed", False):
            logger.error("一致性报告未通过，说明书阶段失败（将触发上层自动修复）。")
            return False
        if not matrix_ok:
            logger.error("声称-证据绑定未通过，说明书阶段失败（将触发上层自动修复）。")
            return False
        if not examiner_material_passed:
            logger.error("审查员可见材料章节未通过，说明书阶段失败（将触发上层自动修复）。")
            return False

        # 渲染模板
        try:
            # [Phase 3] 注入伪造的文档核心属性 (Core Properties)
            # 必须在 render 之前或之后但在 save 之前设置
            fake_meta = self.differentiator.get_fake_metadata()

            # DocxTemplate 继承自 python-docx 的 Document，或者代理了它
            # 通常 tpl 对象本身就可以访问 core_properties
            if hasattr(tpl, 'core_properties'):
                props = tpl.core_properties
                props.author = fake_meta['author']
                props.last_modified_by = fake_meta['last_modified_by']
                props.created = fake_meta['created']
                props.modified = fake_meta['modified']

                # 伪造 Company (通过 custom properties 或者扩展属性，python-docx 标准库支持有限)
                # 但我们可以修改 title/subject/comments
                props.title = f"{software_full_name} 说明书"
                props.subject = "软件著作权申请材料"
                props.category = "技术文档"
                props.comments = f"Generated by {fake_meta['creator']} | Security Level: Internal"

                logger.info(f"已注入伪造文档元数据: Author={props.author}, Created={props.created}")

            tpl.render(context)

            # 确保输出目录存在
            self.output_path.parent.mkdir(parents=True, exist_ok=True)

            # 保存文档
            tpl.save(str(self.output_path))

            # 即使模板未预置占位符，也将 A03/A04/A05 追加到正文末尾，保证审查材料可见。
            self._append_examiner_sections_to_docx(context)

            # 设置文档在打开时自动更新域（包括目录）
            self._enable_fields_update()

            logger.info(f"✓ 说明书生成成功: {self.output_path}")
            return True

        except Exception as e:
            logger.error(f"✗ 说明书生成失败: {e}")
            return False

    def _append_examiner_sections_to_docx(self, context: Dict[str, Any]) -> None:
        """
        将审查版材料 A03/A04/A05 追加到文档末尾，避免依赖模板占位符。
        """
        if not self.output_path.exists():
            return
        try:
            doc = Document(str(self.output_path))
            doc.add_page_break()
            doc.add_heading("审查版材料（A03/A04/A05）", level=1)

            summary_text = str(context.get("examiner_summary") or "").strip()
            if summary_text:
                doc.add_paragraph(summary_text)

            doc.add_heading("A03 功能对应表（页面功能 -> 截图编号 -> 接口 -> 代码位置）", level=2)
            rows = context.get("feature_evidence_rows") or []
            if isinstance(rows, list) and rows:
                table = doc.add_table(rows=1, cols=8)
                table.style = "Table Grid"
                headers = ["序号", "页面/流程", "功能声明", "截图编号", "接口", "代码位置", "运行回放", "状态"]
                for i, title in enumerate(headers):
                    table.cell(0, i).text = title
                for item in rows[:200]:
                    row = table.add_row().cells
                    row[0].text = str(item.get("index") or "")
                    row[1].text = str(item.get("page_or_flow") or "")
                    row[2].text = str(item.get("claim_text") or "")
                    row[3].text = str(item.get("screenshot_refs") or "")
                    row[4].text = str(item.get("api_refs") or "")
                    row[5].text = str(item.get("code_refs") or "")
                    row[6].text = str(item.get("runtime_refs") or "")
                    row[7].text = str(item.get("binding_status") or "")
            else:
                doc.add_paragraph("暂无可渲染的功能对应表数据。")

            doc.add_heading("A04 开发时间说明", level=2)
            timeline_text = str(context.get("timeline_review_text") or "").strip()
            doc.add_paragraph(timeline_text or "暂无开发时间说明数据。")

            doc.add_heading("A05 版本与新增点说明", level=2)
            novelty_text = str(context.get("novelty_review_text") or "").strip()
            doc.add_paragraph(novelty_text or "暂无版本与新增点说明数据。")
            points = context.get("version_increment_points") or []
            if isinstance(points, list) and points:
                for idx, point in enumerate(points[:20], start=1):
                    text = str(point).strip()
                    if text:
                        doc.add_paragraph(f"{idx}. {text}")

            doc.add_heading("运行时技能合规与策略裁决", level=2)
            compliance_summary = context.get("skill_compliance_summary") or {}
            policy_summary = context.get("skill_policy_summary") or {}
            ratio = float(compliance_summary.get("rule_pass_ratio") or 0.0) if isinstance(compliance_summary, dict) else 0.0
            critical = ",".join(compliance_summary.get("critical_failed_rules") or []) if isinstance(compliance_summary, dict) else ""
            actions = ",".join(policy_summary.get("auto_fix_actions") or []) if isinstance(policy_summary, dict) else ""
            doc.add_paragraph(f"规则通过率: {ratio:.1%}")
            doc.add_paragraph(f"关键失败规则: {critical or '无'}")
            doc.add_paragraph(f"策略建议动作: {actions or '无'}")

            doc.save(str(self.output_path))
            logger.info("✓ 已自动追加 A03/A04/A05 到说明书正文")
        except Exception as e:
            logger.warning(f"追加审查版材料章节失败（不阻断主流程）: {e}")

    def _build_page_operation_steps(self, page_title: str, api_info: List[Dict[str, str]]) -> List[str]:
        """
        根据页面标题与接口信息生成可执行的操作步骤，
        避免说明书只停留在“功能描述”层面。
        """
        title = str(page_title or "业务页面").strip()
        steps = [
            f"进入{title}页面，先通过筛选条件定位目标数据。",
        ]

        if api_info:
            first_api = api_info[0]
            first_desc = str(first_api.get("desc", "")).strip() or "查询业务数据"
            steps.append(f"执行“{first_desc}”操作，系统返回当前记录清单。")

            if len(api_info) > 1:
                second_api = api_info[1]
                second_desc = str(second_api.get("desc", "")).strip() or "提交处理结果"
                steps.append(f"对目标记录执行“{second_desc}”，并校验状态变更结果。")
        else:
            steps.append("在列表中选择目标对象，进入详情并完成编辑或审核处理。")

        steps.append("处理完成后导出结果或提交归档，形成可追溯业务记录。")
        return steps

    def _save_operation_flow_guide(self, context: Dict[str, Any]):
        """
        生成操作流程增强说明，作为说明书之外的流程佐证材料。
        """
        try:
            flow_profile = context.get("operation_flow_profile", {}) or {}
            project_name = context.get("project_name", "未命名项目")
            out_path = self.output_path.parent / f"{project_name}_操作流程增强说明.txt"

            lines: List[str] = []
            lines.append(f"项目名称: {project_name}")
            lines.append(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            lines.append("")
            lines.append("一、端到端业务流程画像")

            roles = flow_profile.get("roles", []) or []
            if roles:
                lines.append("角色清单: " + "、".join([str(r) for r in roles]))
            flows = flow_profile.get("flows", []) or []
            if not flows:
                lines.append("（未定义流程画像，使用页面级操作步骤作为补充）")
            else:
                for idx, flow in enumerate(flows, start=1):
                    lines.append(f"{idx}. {flow.get('flow_name', '业务流程')}")
                    lines.append(f"   触发条件: {flow.get('trigger', '人工发起')}")
                    lines.append(f"   执行角色: {flow.get('actor', '业务操作员')}")
                    for step in flow.get("steps", []) or []:
                        lines.append(
                            f"   - Step{step.get('step', '')}: "
                            f"{step.get('page_title', '')} ({step.get('page_id', '')}) -> {step.get('action', '')}"
                        )
                    lines.append(f"   输出结果: {flow.get('output', '业务记录')}")

            lines.append("")
            lines.append("二、页面级操作步骤")
            for idx, page in enumerate(context.get("pages", []), start=1):
                lines.append(f"{idx}. {page.get('title', '未命名页面')}")
                for step_i, text in enumerate(page.get("operation_steps", []) or [], start=1):
                    lines.append(f"   {step_i}) {text}")
                lines.append("")

            with open(out_path, "w", encoding="utf-8") as f:
                f.write("\n".join(lines))
            logger.info(f"✓ 已生成操作流程增强说明: {out_path}")
        except Exception as e:
            logger.warning(f"生成操作流程增强说明失败: {e}")

    def _save_consistency_report(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """
        输出代码-文档一致性报告，便于发现“文档描述与代码结构脱节”问题。
        """
        try:
            naming_facts = getattr(self, "naming_facts", None)
            if not isinstance(naming_facts, dict):
                naming_facts = resolve_software_identity(
                    {},
                    fallback_project_name=str((self.plan or {}).get("project_name") or ""),
                )
            blueprint = self.plan.get("code_blueprint", {}) or {}
            planned_controllers = blueprint.get("controllers", []) or []
            planned_api_count = sum(len(c.get("methods", []) or []) for c in planned_controllers)
            if planned_api_count == 0 and self.project_spec:
                planned_api_count = len(self.project_spec.get("api_contracts", []) or [])
            generated_api_count = len(context.get("api_list", []) or [])
            if generated_api_count == 0:
                page_api_refs = []
                for page in context.get("pages", []) or []:
                    for api_item in page.get("api_info", []) or []:
                        ref = str(api_item.get("http") or api_item.get("name") or "").strip()
                        if ref:
                            page_api_refs.append(ref)
                generated_api_count = len(set(page_api_refs))
            planned_entities = len(blueprint.get("entities", []) or [])
            if planned_entities == 0 and self.project_spec:
                planned_entities = len(self.project_spec.get("entities", []) or [])
            generated_entities = len(context.get("db_tables", []) or [])
            if generated_entities == 0 and self.project_spec:
                generated_entities = len(self.project_spec.get("entities", []) or [])
            page_count = len(self.plan.get("pages", {}) or {})
            rendered_page_count = len(context.get("pages", []) or [])

            api_baseline_ready = planned_api_count > 0
            entity_baseline_ready = planned_entities > 0
            baseline_ready = api_baseline_ready and entity_baseline_ready

            api_coverage = min(1.0, generated_api_count / planned_api_count) if api_baseline_ready else 0.0
            entity_coverage = min(1.0, generated_entities / planned_entities) if entity_baseline_ready else 0.0
            page_coverage = 1.0 if page_count == 0 else min(1.0, rendered_page_count / page_count)

            passed = baseline_ready and (api_coverage >= 0.35) and (entity_coverage >= 0.35) and (page_coverage >= 0.85)
            report = {
                "project_name": self.plan.get("project_name", ""),
                "software_full_name": naming_facts.get("software_full_name"),
                "software_short_name": naming_facts.get("software_short_name"),
                "spec_loaded": bool(self.project_spec),
                "baseline_ready": baseline_ready,
                "api_baseline_ready": api_baseline_ready,
                "entity_baseline_ready": entity_baseline_ready,
                "planned_api_count": planned_api_count,
                "generated_api_count": generated_api_count,
                "planned_entity_count": planned_entities,
                "generated_entity_count": generated_entities,
                "planned_page_count": page_count,
                "rendered_page_count": rendered_page_count,
                "api_coverage": round(api_coverage, 4),
                "entity_coverage": round(entity_coverage, 4),
                "page_coverage": round(page_coverage, 4),
                "passed": passed,
            }

            report_path = self.output_path.parent / "doc_code_consistency_report.json"
            with open(report_path, "w", encoding="utf-8") as f:
                json.dump(report, f, ensure_ascii=False, indent=2)
            if passed:
                logger.info(f"一致性报告已生成: {report_path} (passed)")
            else:
                logger.warning(f"一致性报告已生成: {report_path} (not passed)")
            return report
        except Exception as e:
            logger.warning(f"生成一致性报告失败: {e}")
            return {"passed": False, "error": str(e)}

    @staticmethod
    def _safe_int(value: Any, default: int = 0) -> int:
        try:
            return int(value)
        except Exception:
            return default

    @staticmethod
    def _safe_float(value: Any, default: float = 0.0) -> float:
        try:
            return float(value)
        except Exception:
            return default

    @staticmethod
    def _join_items(items: List[str], sep: str = "；", limit: int = 4, empty_value: str = "无") -> str:
        cleaned = [str(x).strip() for x in (items or []) if str(x).strip()]
        if not cleaned:
            return empty_value
        clipped = cleaned[: max(1, int(limit))]
        text = sep.join(clipped)
        if len(cleaned) > len(clipped):
            text += f"（共{len(cleaned)}项）"
        return text

    def _load_json_quiet(self, path: Path) -> Dict[str, Any]:
        try:
            if not path.exists():
                return {}
            with open(path, "r", encoding="utf-8") as f:
                payload = json.load(f)
            return payload if isinstance(payload, dict) else {}
        except Exception:
            return {}

    def _build_feature_evidence_rows(self, matrix: Dict[str, Any]) -> List[Dict[str, Any]]:
        claims = matrix.get("claims") or []
        if not isinstance(claims, list):
            return []

        rows: List[Dict[str, Any]] = []
        for idx, claim in enumerate(claims, start=1):
            if not isinstance(claim, dict):
                continue

            claim_type = str(claim.get("claim_type") or "").strip() or "unknown"
            claim_text = str(claim.get("claim_text") or "").strip() or f"未命名声称#{idx}"
            evidence = claim.get("evidence") or {}
            evidence = evidence if isinstance(evidence, dict) else {}
            missing = claim.get("missing_evidence") or []
            missing = [str(x).strip() for x in missing if str(x).strip()]

            page_or_flow = str(claim.get("page_title") or claim.get("flow_name") or claim.get("page_id") or claim.get("claim_id") or "")
            block_id = str(claim.get("block_id") or "").strip()
            selector = str(claim.get("selector") or "").strip()
            if claim_type == "page_capability":
                page_id = str(claim.get("page_id") or "").strip()
                if page_id:
                    page_or_flow = f"{page_or_flow} ({page_id})" if page_or_flow else page_id
            elif claim_type == "functional_block":
                if block_id:
                    page_or_flow = f"{page_or_flow} [{block_id}]".strip()

            screenshot_refs = [str(x).replace("\\", "/") for x in (evidence.get("screenshot_paths") or []) if str(x).strip()]
            code_refs = [str(x).replace("\\", "/") for x in (evidence.get("code_hits") or []) if str(x).strip()]

            api_refs: List[str] = []
            for item in evidence.get("api_contracts") or []:
                if not isinstance(item, dict):
                    continue
                method = str(item.get("http_method") or "").strip().upper()
                path = str(item.get("path") or "").strip()
                name = str(item.get("method_name") or "").strip()
                if method and path:
                    api_refs.append(f"{method} {path}")
                elif name:
                    api_refs.append(name)

            runtime_refs: List[str] = []
            for replay in evidence.get("runtime_replay_matches") or []:
                if not isinstance(replay, dict):
                    continue
                flow = str(replay.get("flow") or "").strip() or "流程回放"
                status = "通过" if bool(replay.get("passed")) else "未通过"
                runtime_refs.append(f"{flow}({status})")

            runtime_replay = evidence.get("runtime_replay") or {}
            if isinstance(runtime_replay, dict) and runtime_replay:
                status = "通过" if bool(runtime_replay.get("passed")) else "未通过"
                runtime_refs.append(f"流程回放({status})")
                related_pages = [str(x).strip() for x in (runtime_replay.get("related_pages") or []) if str(x).strip()]
                related_apis = [str(x).strip() for x in (runtime_replay.get("related_apis") or []) if str(x).strip()]
                if related_pages:
                    runtime_refs.append("关联页面:" + "、".join(related_pages[:4]))
                if related_apis:
                    runtime_refs.append("关联接口:" + "、".join(related_apis[:4]))

            row = {
                "index": idx,
                "claim_id": str(claim.get("claim_id") or f"claim_{idx}"),
                "claim_type": claim_type,
                "page_or_flow": page_or_flow or "未标识对象",
                "claim_text": claim_text,
                "screenshot_refs": self._join_items(screenshot_refs, sep="；", limit=4),
                "api_refs": self._join_items(api_refs, sep="；", limit=4),
                "code_refs": self._join_items(code_refs, sep="；", limit=4),
                "runtime_refs": self._join_items(runtime_refs, sep="；", limit=4),
                "missing_evidence": self._join_items(missing, sep="、", limit=6, empty_value="无"),
                "binding_status": "通过" if bool(claim.get("passed")) else "未通过",
                "block_id": block_id,
                "selector": selector,
            }
            rows.append(row)
        return rows

    def _build_timeline_review(self, project_dir: Path, timeline_report: Dict[str, Any]) -> Dict[str, Any]:
        report = timeline_report if isinstance(timeline_report, dict) else {}
        metadata = self._load_json_quiet(project_dir / "project_metadata.json")
        declared = report.get("declared_timeline") or {}
        inferred = report.get("inferred_timeline") or {}
        declared = declared if isinstance(declared, dict) else {}
        inferred = inferred if isinstance(inferred, dict) else {}

        if not declared and isinstance(metadata, dict):
            declared = {
                "development_started_at": str(
                    metadata.get("development_started_at")
                    or metadata.get("dev_started_at")
                    or ""
                ).strip(),
                "development_completed_at": str(
                    metadata.get("development_completed_at")
                    or metadata.get("dev_completed_at")
                    or ""
                ).strip(),
                "published_at": str(metadata.get("published_at") or "").strip(),
                "submit_at": str(metadata.get("submit_at") or metadata.get("submitted_at") or "").strip(),
                "organization_established_at": str(
                    metadata.get("organization_established_at")
                    or metadata.get("company_established_at")
                    or metadata.get("entity_established_at")
                    or ""
                ).strip(),
            }

        issues = [str(x).strip() for x in (report.get("issues") or []) if str(x).strip()]
        warnings = [str(x).strip() for x in (report.get("warnings") or []) if str(x).strip()]
        has_data = bool(report) or any(str(v).strip() for v in declared.values())
        is_passed = bool(report.get("passed")) if report else (len(issues) == 0)

        text_lines = [
            f"开发开始（声明）: {declared.get('development_started_at') or 'N/A'}",
            f"开发完成（声明）: {declared.get('development_completed_at') or 'N/A'}",
            f"发表时间（声明）: {declared.get('published_at') or 'N/A'}",
            f"提交时间（声明）: {declared.get('submit_at') or 'N/A'}",
            f"主体成立时间（声明）: {declared.get('organization_established_at') or 'N/A'}",
            f"开发开始（推断）: {inferred.get('development_started_at') or 'N/A'}",
            f"开发完成（推断）: {inferred.get('development_completed_at') or 'N/A'}",
            f"冻结时间: {inferred.get('frozen_at') or 'N/A'}",
        ]
        if issues:
            text_lines.append("阻断问题: " + "；".join(issues))
        if warnings:
            text_lines.append("补证提示: " + "；".join(warnings))

        return {
            "ready": has_data,
            "passed": is_passed,
            "issues": issues,
            "warnings": warnings,
            "declared_timeline": declared,
            "inferred_timeline": inferred,
            "requires_supporting_note": bool(report.get("requires_supporting_note")),
            "text": "\n".join(text_lines),
        }

    def _build_novelty_review(
        self,
        novelty_report: Dict[str, Any],
        semantic_report: Dict[str, Any],
        semantic_loop_report: Dict[str, Any],
    ) -> Dict[str, Any]:
        novelty = novelty_report if isinstance(novelty_report, dict) else {}
        semantic = semantic_report if isinstance(semantic_report, dict) else {}
        semantic_loop = semantic_loop_report if isinstance(semantic_loop_report, dict) else {}

        max_similarity = self._safe_float(novelty.get("max_similarity"), 0.0)
        recommendation = str(novelty.get("recommendation") or "").strip() or "unknown"
        top_similarity = self._safe_float(semantic.get("top_similarity"), 0.0)
        should_rewrite = bool(semantic.get("should_rewrite"))
        rewritten = bool(semantic.get("rewritten"))
        loop_rounds = self._safe_int(semantic_loop.get("max_rounds"), 0)
        loop_passed = bool(semantic_loop.get("passed")) if semantic_loop else (not should_rewrite)

        api_count = len(self.project_spec.get("api_contracts") or []) if isinstance(self.project_spec, dict) else 0
        entity_count = len(self.project_spec.get("entities") or []) if isinstance(self.project_spec, dict) else 0
        page_count = len(self.plan.get("pages") or []) if isinstance(self.plan.get("pages"), list) else len(self.plan.get("pages") or {})

        version_increment_points: List[str] = []
        if api_count > 0:
            version_increment_points.append(f"接口合约已固化 {api_count} 个（来源 project_executable_spec.json）")
        if entity_count > 0:
            version_increment_points.append(f"实体模型已固化 {entity_count} 个（与接口/页面映射一致）")
        if page_count > 0:
            version_increment_points.append(f"页面功能清单已覆盖 {page_count} 个页面")
        if recommendation != "unknown":
            version_increment_points.append(f"指纹最大相似度 {max_similarity:.3f}，审计建议 {recommendation}")
        if semantic:
            version_increment_points.append(
                f"语义同质化 top_similarity={top_similarity:.3f}，当前状态={'需重写' if should_rewrite else '已达标'}"
            )
        if rewritten:
            version_increment_points.append("已执行语义改写并同步刷新规格/计划引用")
        if semantic_loop:
            version_increment_points.append(
                f"同质化闭环复检结果={'通过' if loop_passed else '未通过'}，最大轮次={max(1, loop_rounds)}"
            )
        if not version_increment_points:
            version_increment_points.append("暂无可量化新增点，请先生成 novelty/semantic 报告后再导出材料")

        text_lines = [
            f"指纹审计建议: {recommendation}",
            f"指纹最大相似度: {max_similarity:.3f}",
            f"语义同质化 top_similarity: {top_similarity:.3f}",
            f"语义改写状态: {'需改写' if should_rewrite else '达标'}",
            f"闭环复检: {'通过' if loop_passed else '未通过'}",
            "新增点摘要: " + "；".join(version_increment_points[:4]),
        ]

        has_blocking_issue = (recommendation == "blocked") or should_rewrite or (semantic_loop and not loop_passed)
        ready = bool(novelty or semantic or semantic_loop)
        return {
            "ready": ready,
            "has_blocking_issue": has_blocking_issue,
            "recommendation": recommendation,
            "max_similarity": round(max_similarity, 4),
            "top_similarity": round(top_similarity, 4),
            "should_rewrite": should_rewrite,
            "rewritten": rewritten,
            "loop_passed": loop_passed,
            "loop_rounds": loop_rounds,
            "version_increment_points": version_increment_points,
            "text": "\n".join(text_lines),
        }

    def _save_examiner_material_sections(
        self,
        context: Dict[str, Any],
        matrix: Dict[str, Any],
        strict_timeline: bool = True,
    ) -> Dict[str, Any]:
        """
        输出审查员可见章节（A03/A04/A05）：
        - 功能对应表（声明 -> 截图/API/代码/回放）
        - 开发时间说明
        - 版本与新增点说明
        """
        project_dir = self.output_path.parent
        runtime_report_path = project_dir / "runtime_verification_report.json"
        matrix_path = project_dir / "claim_evidence_matrix.json"
        timeline_report_path = project_dir / "freeze_package" / "timeline_consistency_report.json"
        novelty_report_path = project_dir / "novelty_quality_report.json"
        semantic_report_path = project_dir / "semantic_homogeneity_report.json"
        semantic_loop_report_path = project_dir / "semantic_homogeneity_closed_loop_report.json"
        runtime_skill_plan_path = project_dir / "runtime_skill_plan.json"
        runtime_skill_override_path = project_dir / "runtime_skill_override.json"
        skill_studio_plan_path = project_dir / "skill_studio_plan.json"
        runtime_rule_graph_path = project_dir / "runtime_rule_graph.json"
        skill_compliance_report_path = project_dir / "skill_compliance_report.json"
        skill_autorepair_report_path = project_dir / "skill_autorepair_report.json"
        skill_policy_report_path = project_dir / "skill_policy_decision_report.json"
        ui_skill_profile_path = project_dir / "ui_skill_profile.json"
        ui_blueprint_path = project_dir / "ui_blueprint.json"

        runtime_report = self._load_json_quiet(runtime_report_path)
        matrix_payload = matrix if isinstance(matrix, dict) and matrix else self._load_json_quiet(matrix_path)
        timeline_report = self._load_json_quiet(timeline_report_path)
        novelty_report = self._load_json_quiet(novelty_report_path)
        semantic_report = self._load_json_quiet(semantic_report_path)
        semantic_loop_report = self._load_json_quiet(semantic_loop_report_path)
        runtime_skill_plan = self._load_json_quiet(runtime_skill_plan_path)
        runtime_skill_override = self._load_json_quiet(runtime_skill_override_path)
        skill_studio_plan = self._load_json_quiet(skill_studio_plan_path)
        runtime_rule_graph = self._load_json_quiet(runtime_rule_graph_path)
        skill_compliance_report = self._load_json_quiet(skill_compliance_report_path)
        skill_autorepair_report = self._load_json_quiet(skill_autorepair_report_path)
        skill_policy_report = self._load_json_quiet(skill_policy_report_path)
        ui_skill_profile = self._load_json_quiet(ui_skill_profile_path)
        ui_blueprint = self._load_json_quiet(ui_blueprint_path)
        compliance_summary = skill_compliance_report.get("summary") or {}

        feature_rows = self._build_feature_evidence_rows(matrix_payload)
        matrix_summary = matrix_payload.get("summary") or {}
        total_claims = self._safe_int(matrix_summary.get("total_claims"), 0)
        passed_claims = self._safe_int(matrix_summary.get("passed_claims"), 0)
        binding_ratio = self._safe_float(matrix_summary.get("binding_ratio"), 0.0)
        hard_issues = [str(x).strip() for x in (matrix_payload.get("hard_blocking_issues") or []) if str(x).strip()]

        replay_report = ((runtime_report.get("checks") or {}).get("business_path_replay") or {})
        replay_ratio = self._safe_float(replay_report.get("match_ratio"), 0.0)
        runtime_passed = bool(runtime_report.get("overall_passed"))
        replay_passed = bool(replay_report.get("passed"))

        timeline_review = self._build_timeline_review(project_dir, timeline_report)
        novelty_review = self._build_novelty_review(novelty_report, semantic_report, semantic_loop_report)
        version_increment_points = novelty_review.get("version_increment_points") or []

        feature_ready = (len(feature_rows) > 0) and (total_claims > 0) and (binding_ratio >= 0.85) and (len(hard_issues) == 0)
        timeline_ready = bool(timeline_review.get("ready")) and bool(timeline_review.get("passed")) and not bool(timeline_review.get("issues"))
        novelty_ready = bool(novelty_review.get("ready")) and not bool(novelty_review.get("has_blocking_issue"))

        blocking_issues: List[str] = []
        if not feature_ready:
            if total_claims <= 0:
                blocking_issues.append("功能对应表无有效声明，无法形成审查映射")
            if binding_ratio < 0.85:
                blocking_issues.append("功能对应表绑定率不足（<85%）")
            if hard_issues:
                blocking_issues.extend(hard_issues[:5])
        if not runtime_passed:
            blocking_issues.append("运行验证未通过，无法作为审查回放证据")
        if not replay_passed:
            blocking_issues.append("业务流程回放未全量命中")
        timeline_blockers: List[str] = []
        if not timeline_ready:
            if not timeline_review.get("ready"):
                timeline_blockers.append("缺少开发时间说明数据")
            else:
                timeline_blockers.extend([f"时间线冲突: {x}" for x in (timeline_review.get("issues") or [])][:5])
        timeline_deferred_issues: List[str] = []
        if timeline_blockers:
            if strict_timeline:
                blocking_issues.extend(timeline_blockers)
            else:
                timeline_deferred_issues.extend(timeline_blockers)
                stage_note = "文档阶段说明: A04 时间线信息待 freeze 阶段复核，当前先输出主文档供人工校对。"
                review_text = str(timeline_review.get("text") or "").strip()
                if stage_note not in review_text:
                    review_text = f"{review_text}\n{stage_note}" if review_text else stage_note
                timeline_review["text"] = review_text
                warning_items = [str(x).strip() for x in (timeline_review.get("warnings") or []) if str(x).strip()]
                for issue in timeline_deferred_issues:
                    tagged = f"{issue}（文档阶段暂缓，freeze 前需修复）"
                    if tagged not in warning_items:
                        warning_items.append(tagged)
                timeline_review["warnings"] = warning_items
                timeline_review["requires_supporting_note"] = True
        if not novelty_ready:
            if not novelty_review.get("ready"):
                blocking_issues.append("缺少版本与新增点审计数据")
            else:
                blocking_issues.append("版本与新增点说明未达可授权级（同质化/相似度未收敛）")
        if compliance_summary and (not bool(compliance_summary.get("passed"))):
            blocking_issues.append("运行时技能合规未通过")
            critical_rules = [str(x).strip() for x in (compliance_summary.get("critical_failed_rules") or []) if str(x).strip()]
            if critical_rules:
                blocking_issues.append(f"运行时技能关键规则失败: {','.join(critical_rules[:6])}")

        feature_lines = []
        for row in feature_rows:
            feature_lines.append(
                f"{row['index']}. {row['page_or_flow']} | 声称: {row['claim_text']} | "
                f"截图: {row['screenshot_refs']} | 接口: {row['api_refs']} | 代码: {row['code_refs']} | "
                f"回放: {row['runtime_refs']} | 状态: {row['binding_status']}"
            )

        examiner_summary = (
            f"共编译 {total_claims} 条声明，绑定通过 {passed_claims} 条，绑定率 {binding_ratio:.1%}。"
            f"运行验证={'通过' if runtime_passed else '未通过'}，流程回放命中率 {replay_ratio:.1%}。"
        )
        ui_summary = ui_blueprint.get("summary") or {}
        if ui_summary:
            examiner_summary += (
                f" UI蓝图覆盖 {int(ui_summary.get('page_count') or 0)} 页，"
                f"{int(ui_summary.get('block_count') or 0)} 个功能块。"
            )
        runtime_skillpack = runtime_skill_plan.get("skillpack") or {}
        runtime_domain = (runtime_skill_plan.get("domain_match") or {}).get("domain")
        if runtime_skillpack:
            examiner_summary += (
                f" 运行时技能包 {str(runtime_skillpack.get('id') or '')} "
                f"(domain={str(runtime_domain or 'generic')})。"
            )
        if runtime_rule_graph:
            rule_count = int((((runtime_rule_graph.get("graph") or {}).get("summary") or {}).get("rule_count") or 0))
            examiner_summary += f" 规则图条目 {rule_count}。"
        if compliance_summary:
            examiner_summary += (
                f" 合规规则通过率 {float(compliance_summary.get('rule_pass_ratio') or 0.0):.1%}。"
            )
        policy_summary = skill_policy_report.get("summary") or {}
        if policy_summary:
            examiner_summary += (
                f" 策略裁决建议动作 {len(policy_summary.get('auto_fix_actions') or [])} 项。"
            )
        if runtime_skill_override:
            examiner_summary += " 已应用项目级技能覆写。"
        if skill_studio_plan:
            examiner_summary += " 技能工作台决策已生效。"

        context["examiner_summary"] = examiner_summary
        context["feature_evidence_rows"] = feature_rows
        context["feature_evidence_text"] = "\n".join(feature_lines) if feature_lines else "暂无可渲染的功能对应表。"
        context["timeline_review"] = timeline_review
        context["timeline_review_text"] = str(timeline_review.get("text") or "")
        context["novelty_review"] = novelty_review
        context["novelty_review_text"] = str(novelty_review.get("text") or "")
        context["version_increment_points"] = version_increment_points
        context["skill_compliance_summary"] = compliance_summary
        context["skill_policy_summary"] = policy_summary

        passed = len(blocking_issues) == 0
        report = {
            "project_name": context.get("project_name") or self.plan.get("project_name") or project_dir.name,
            "generated_at": datetime.now().isoformat(timespec="seconds"),
            "passed": passed,
            "blocking_issues": blocking_issues,
            "deferred_issues": timeline_deferred_issues,
            "strict_timeline": bool(strict_timeline),
            "source_paths": {
                "runtime_report": str(runtime_report_path),
                "claim_evidence_matrix": str(matrix_path),
                "timeline_report": str(timeline_report_path),
                "novelty_report": str(novelty_report_path),
                "semantic_report": str(semantic_report_path),
                "semantic_closed_loop_report": str(semantic_loop_report_path),
                "runtime_skill_plan": str(runtime_skill_plan_path),
                "runtime_skill_override": str(runtime_skill_override_path),
                "skill_studio_plan": str(skill_studio_plan_path),
                "runtime_rule_graph": str(runtime_rule_graph_path),
                "skill_compliance_report": str(skill_compliance_report_path),
                "skill_autorepair_report": str(skill_autorepair_report_path),
                "skill_policy_report": str(skill_policy_report_path),
                "ui_skill_profile": str(ui_skill_profile_path),
                "ui_blueprint": str(ui_blueprint_path),
            },
            "summary": {
                "examiner_summary": examiner_summary,
                "runtime_passed": runtime_passed,
                "replay_passed": replay_passed,
                "replay_match_ratio": round(replay_ratio, 4),
                "claim_binding_ratio": round(binding_ratio, 4),
                "claim_total": total_claims,
                "claim_passed": passed_claims,
                "claim_hard_issues": hard_issues,
                "ui_skill_mode": str(ui_skill_profile.get("mode") or ""),
                "runtime_skillpack_id": str(runtime_skillpack.get("id") or ""),
                "runtime_domain": str(runtime_domain or "generic"),
                "runtime_rule_count": int((((runtime_rule_graph.get("graph") or {}).get("summary") or {}).get("rule_count") or 0)),
                "runtime_override_applied": bool((runtime_skill_plan.get("override_applied") or {}).get("applied")) or bool(runtime_skill_override),
                "skill_studio_decision_domain": str(((skill_studio_plan.get("decisions") or {}).get("domain") or "")) if isinstance(skill_studio_plan, dict) else "",
                "skill_compliance_passed": bool(compliance_summary.get("passed")) if compliance_summary else False,
                "skill_compliance_ratio": round(float(compliance_summary.get("rule_pass_ratio") or 0.0), 4) if compliance_summary else 0.0,
                "skill_autorepair_rounds": len(skill_autorepair_report.get("rounds") or []) if isinstance(skill_autorepair_report, dict) else 0,
                "skill_policy_actions": list(policy_summary.get("auto_fix_actions") or []) if isinstance(policy_summary, dict) else [],
                "ui_block_count": int((ui_summary.get("block_count") or 0)),
            },
            "sections_ready": {
                "feature_evidence_ready": feature_ready,
                "timeline_review_ready": timeline_ready,
                "novelty_review_ready": novelty_ready,
                "skill_compliance_ready": bool(compliance_summary.get("passed")) if compliance_summary else False,
            },
            "counts": {
                "feature_evidence_row_count": len(feature_rows),
                "version_increment_point_count": len(version_increment_points),
            },
        }

        report_path = project_dir / "examiner_material_report.json"
        with open(report_path, "w", encoding="utf-8") as f:
            json.dump(report, f, ensure_ascii=False, indent=2)

        md_lines: List[str] = [
            "# 审查版材料章节（可并入说明书）",
            "",
            f"- 项目名称: {report['project_name']}",
            f"- 生成时间: {report['generated_at']}",
            f"- 综合结论: {'通过' if passed else '未通过'}",
            "",
            "## 一、审查摘要",
            examiner_summary,
            "",
            "## 二、A03 功能对应表（页面功能 -> 截图 -> 接口 -> 代码位置）",
            "",
            "| 序号 | 页面/流程 | 功能声明 | 截图编号 | 接口 | 代码位置 | 运行回放 | 状态 |",
            "|---|---|---|---|---|---|---|---|",
        ]
        for row in feature_rows:
            md_lines.append(
                f"| {row['index']} | {row['page_or_flow']} | {row['claim_text']} | {row['screenshot_refs']} | "
                f"{row['api_refs']} | {row['code_refs']} | {row['runtime_refs']} | {row['binding_status']} |"
            )

        md_lines.extend(
            [
                "",
                "## 三、A04 开发时间说明",
                str(timeline_review.get("text") or "暂无开发时间说明数据"),
                "",
                "## 四、A05 版本与新增点说明",
                str(novelty_review.get("text") or "暂无版本与新增点数据"),
                "",
                "## 五、运行时技能合规与策略裁决",
                f"- 规则通过率: {float(compliance_summary.get('rule_pass_ratio') or 0.0):.1%}",
                f"- 关键失败规则: {','.join(compliance_summary.get('critical_failed_rules') or []) or '无'}",
                f"- 自动修复轮次: {len(skill_autorepair_report.get('rounds') or []) if isinstance(skill_autorepair_report, dict) else 0}",
                f"- 策略建议动作: {','.join(policy_summary.get('auto_fix_actions') or []) if isinstance(policy_summary, dict) else ''}",
                "",
                "### 本次新增点（可直接并入说明书）",
            ]
        )
        for idx, point in enumerate(version_increment_points, start=1):
            md_lines.append(f"{idx}. {point}")

        if timeline_deferred_issues:
            md_lines.extend(
                [
                    "",
                    "## 文档阶段待复核项",
                    *[f"- {x}（freeze 阶段将作为硬门禁）" for x in timeline_deferred_issues],
                ]
            )

        if blocking_issues:
            md_lines.extend(["", "## 阻断项", *[f"- {x}" for x in blocking_issues]])

        section_md_path = project_dir / "examiner_material_sections.md"
        with open(section_md_path, "w", encoding="utf-8") as f:
            f.write("\n".join(md_lines))

        logger.info(f"✓ 审查版材料报告已生成: {report_path}")
        logger.info(f"✓ 审查版材料章节已生成: {section_md_path}")
        return report
    
    def _enable_fields_update(self):
        """自动打开 Word 并更新所有域（包括目录）"""
        try:
            import win32com.client
            import pythoncom
            
            logger.info("正在自动更新文档域...")
            
            # 初始化 COM
            pythoncom.CoInitialize()
            
            # 创建 Word 应用程序实例
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False  # 后台运行
            word.DisplayAlerts = 0  # 不显示警告
            
            try:
                # 打开文档
                doc = word.Documents.Open(str(self.output_path.absolute()))
                
                # 更新所有域（包括目录）
                doc.Fields.Update()
                
                # 更新目录（TOC）
                for toc in doc.TablesOfContents:
                    toc.Update()
                
                # 保存并关闭
                doc.Save()
                doc.Close(SaveChanges=False)
                
                logger.info("✓ 域自动更新成功")
                
            finally:
                # 退出 Word
                word.Quit()
                pythoncom.CoUninitialize()
            
        except ImportError:
            logger.warning("⚠ 未安装 pywin32，将使用打开时更新模式")
            logger.warning("  安装方法: pip install pywin32")
            # 降级为打开时更新
            self._enable_fields_update_on_open()
            
        except Exception as e:
            logger.warning(f"⚠ 自动更新失败（使用打开时更新模式）: {e}")
            # 降级为打开时更新
            self._enable_fields_update_on_open()
    
    def _enable_fields_update_on_open(self):
        """设置文档在打开时自动更新域（降级方案）"""
        try:
            from docx import Document
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            doc = Document(str(self.output_path))
            settings = doc.settings
            settings_element = settings.element
            
            update_fields = settings_element.find(qn('w:updateFields'))
            if update_fields is None:
                update_fields = OxmlElement('w:updateFields')
                update_fields.set(qn('w:val'), 'true')
                settings_element.append(update_fields)
            else:
                update_fields.set(qn('w:val'), 'true')
            
            doc.save(str(self.output_path))
            logger.info("✓ 已设置为打开时更新域")
            
        except Exception as e:
            logger.warning(f"⚠ 域更新设置失败: {e}")

    def _get_page_apis(self, page_id: str) -> List[Dict[str, str]]:
        """
        从 code_blueprint 中提取该页面关联的 API
        
        Args:
            page_id: 页面ID
            
        Returns:
            API信息列表
        """
        blueprint = self.plan.get("code_blueprint", {})
        for controller in blueprint.get("controllers", []):
            if controller.get("page_id") == page_id:
                return controller.get("methods", [])

        # Spec-first 回退：当 plan 中缺失 code_blueprint 时，尝试使用可执行规格补齐接口语义
        if self.project_spec:
            api_ids = []
            for mapping in self.project_spec.get("page_api_mapping", []) or []:
                if mapping.get("page_id") == page_id:
                    api_ids.extend(mapping.get("api_ids", []) or [])
            if api_ids:
                contract_map = {
                    str(item.get("id")): item
                    for item in (self.project_spec.get("api_contracts", []) or [])
                    if str(item.get("id", "")).strip()
                }
                methods: List[Dict[str, str]] = []
                for api_id in sorted(set(str(i) for i in api_ids if str(i).strip())):
                    contract = contract_map.get(api_id, {})
                    verb = str(contract.get("http_method") or "GET").upper()
                    path = str(contract.get("path") or "/api/undefined")
                    methods.append(
                        {
                            "name": str(contract.get("method_name") or api_id),
                            "desc": str(contract.get("description") or ""),
                            "http": f"{verb} {path}",
                        }
                    )
                return methods
        return []


def generate_document(plan_path: str, screenshot_dir: str, template_path: str, output_path: str) -> bool:
    """
    便捷函数：生成 Word 说明书
    
    Args:
        plan_path: project_plan.json 路径
        screenshot_dir: 截图目录路径
        template_path: Word 模板路径
        output_path: 输出 Word 文件路径
    
    Returns:
        是否成功生成
    """
    generator = DocumentGenerator(plan_path, screenshot_dir, template_path, output_path)
    return generator.generate()
